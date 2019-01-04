''' Initializing additional libs '''
import os, traceback, logging
import time
import re
from pprint import pformat, pprint
from tabulate import tabulate
from lib.utils import check_output
EXPORT_DISABLE = False
try:
	from docx import Document
	# Import `load_workbook` module from `openpyxl`
	#from openpyxl import load_workbook
	import openpyxl
except ImportError:
	print('[Err] : Modules for export have not loaded. Exporting will be disabled')
	EXPORT_DISABLE = True

def xlsx_export(results, BASE_DIR):
	'''
	Export result of executing command to Excel
		Arg:
			results: List with all data
			BASE_DIR: base directory of script
	'''
	if EXPORT_DISABLE: 
		logging.debug('[Err] : Modules for export have not loaded. Exporting will be disabled')
		return
	print ('Start export report to excel.....')
	logging.info('Start export report to excel.....')
	result_table = []
	 
	for node in results:
		#pprint(node)
		# create a new workbook and select the active worksheet
		workbook = openpyxl.Workbook()
		
		for cmd_index, command in enumerate(node.get(list(node.keys())[0])):
			#pprint(command)
			if command.get('search_type') not in ['complex', 'TextFSM']: #!= 'complex' or command.get('search_type') != 'TextFSM':
				print ('Command: {} was skipped, search type: {} is incompatible'.format(command.get('cmdname'), command.get('search_type')))
				logging.info('Command: {} was skipped, search type: {} is incompatible'.format(command.get('cmdname'), command.get('search_type')))
				continue
			workbook.create_sheet(command.get('cmdname'))
			#worksheet = workbook.get_sheet_by_name(command.get('cmdname'))
			worksheet = workbook[command.get('cmdname')]
			#worksheet = workbook.active
			result_list_cmd = check_output(command, True)
			logging.debug('Result list of command for export to excel: {}'.format(pformat(result_list_cmd)))
			#pprint(result_list_cmd)
			for row_index, row in enumerate(result_list_cmd):
				for coll_index, coll in enumerate(row):
					worksheet.cell(row=row_index+1, column=coll_index+1).value = str(coll)
		workbook.save('{}report/{}_{}.xlsx'.format(BASE_DIR, list(node.keys())[0], time.strftime("%Y%m%d%H%M%S")))
	print ('Export report to excel finished')
	logging.info('Export report to excel finished')
	
def docx_export(results, BASE_DIR):
	'''
	Export result of executing command to Word
		Arg:
			results: List with all data
			BASE_DIR: base directory of script
	'''
	if EXPORT_DISABLE: 
		logging.debug('[Err] : Modules for export have not loaded. Exporting will be disabled')
		return
	print ('Start export report to word.....')
	logging.info('Start export report to word.....')
	for node in results:
		# Load in the wordbook
		if os.path.isfile(BASE_DIR + 'data/protocol_template.docx') :
			document = Document(BASE_DIR + 'data/protocol_template.docx')
		else :
			print('[Err] protocol_template.docx does not exist.\nExport report to word fail.')
			return 
		for table in document.tables:
			if table.cell(0, 0).text.find('#') >= 0:
				for cmd_index, command in enumerate(node.get(list(node.keys())[0])):
					if table.cell(0, 0).text.find(command.get('command')) >= 0:
						#table.cell(1, 0).add_paragraph(command.get('output'), 'a3')
						table.cell(1, 0).text = command.get('output')
						continue
		document.save('{}report/{}_protocol_{}.docx'.format(BASE_DIR, list(node.keys())[0], time.strftime("%Y%m%d%H%M%S")))
	print ('Export report to word finished')
	logging.info('Export report to word finished')	
